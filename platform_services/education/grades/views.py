from platform_services.identity.mixins import TenantQuerySetMixin
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework import viewsets
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from django.utils import timezone


from .models import Grade, GradeHistory, SequenceValidation
from .serializers import GradeHistorySerializer, GradeSerializer, SequenceValidationSerializer
from .services import calculate_student_bulletin


class GradeViewSet(TenantQuerySetMixin, viewsets.ModelViewSet):
    queryset = Grade.objects.select_related('student', 'subject', 'teacher', 'academic_year').all()
    serializer_class = GradeSerializer

    def get_queryset(self):
        queryset = super().get_queryset()
        user = self.request.user
        if user.is_authenticated and user.role == 'enseignant':
            return queryset.filter(teacher__user=user)
        return queryset

    def check_is_locked(self, student, academic_year_id, sequence):
        from platform_services.education.classes.services import CalendarEngine
        org_id = getattr(self.request.user, 'organization_id', None)
        if CalendarEngine.are_grades_locked(timezone.now().date(), organization_id=org_id):
            return True

        # Need to get class from enrollment now
        enrollment = student.enrollments.filter(academic_year_id=academic_year_id).first()
        if not enrollment:
            return False
        school_class = enrollment.school_class

        validation = SequenceValidation.objects.filter(
            school_class=school_class,
            academic_year_id=academic_year_id,
            sequence=sequence
        ).first()
        return validation.is_locked if validation else False

    def perform_create(self, serializer):
        student = serializer.validated_data['student']
        academic_year = serializer.validated_data['academic_year']
        sequence = serializer.validated_data['sequence']
        
        if self.check_is_locked(student, academic_year.id, sequence):
            from rest_framework.exceptions import ValidationError
            raise ValidationError("La saisie est verrouillée pour cette séquence et cette classe.")
            
        serializer.save()

    def perform_update(self, serializer):
        grade = self.get_object()
        student = grade.student
        
        if self.check_is_locked(student, grade.academic_year.id, grade.sequence):
            from rest_framework.exceptions import ValidationError
            raise ValidationError("La saisie est verrouillée pour cette séquence et cette classe.")
            
        new_value = serializer.validated_data.get('value', grade.value)
        reason = serializer.validated_data.get('reason', '')
        
        if new_value != grade.value and not reason:
            from rest_framework.exceptions import ValidationError
            raise ValidationError({"reason": "Un motif est obligatoire pour modifier une note existante."})
            
        if new_value != grade.value:
            GradeHistory.objects.create(
                grade=grade,
                old_value=grade.value,
                new_value=new_value,
                reason=reason,
                changed_by=self.request.user
            )
            
        serializer.save()

    @action(detail=False, methods=['get'])
    def bulletin(self, request):
        student_id = request.query_params.get('student_id')
        academic_year_id = request.query_params.get('academic_year_id')
        sequence = request.query_params.get('sequence')

        if not all([student_id, academic_year_id, sequence]):
            return Response({'error': 'Veuillez fournir student_id, academic_year_id et sequence.'}, status=400)
            
        data = calculate_student_bulletin(student_id, academic_year_id, sequence)
        return Response(data)

    @action(detail=False, methods=['post'])
    def batch(self, request):
        grades_data = request.data.get('grades', [])
        
        created = 0
        updated = 0
        errors = []
        
        # Préchauffer les verrouillages pour optimiser
        locked_cache = {}
        from platform_services.education.students.models import Student
        students_cache = {s.id: s for s in Student.objects.filter(id__in=[g.get('student') for g in grades_data if g.get('student')])}
        
        for g_data in grades_data:
            student_id = g_data.get('student')
            subject_id = g_data.get('subject')
            sequence = g_data.get('sequence')
            academic_year_id = g_data.get('academic_year')
            evaluation_type = g_data.get('evaluation_type')
            new_value = g_data.get('value')
            reason = g_data.get('reason', '')
            
            if not all([student_id, subject_id, sequence, academic_year_id, evaluation_type, new_value is not None]):
                continue
                
            student = students_cache.get(student_id)
            if not student:
                continue

            enrollment = student.enrollments.filter(academic_year_id=academic_year_id).first()
            if not enrollment:
                errors.append(f"Élève {student.last_name} non inscrit pour l'année académique spécifiée.")
                continue
            
            school_class = enrollment.school_class

            cache_key = f"{school_class.id}_{academic_year_id}_{sequence}"
            if cache_key not in locked_cache:
                locked_cache[cache_key] = self.check_is_locked(student, academic_year_id, sequence)
                
            if locked_cache[cache_key]:
                errors.append(f"Élève {student.last_name} bloqué (séquence verrouillée).")
                continue
                
            grade_obj, is_new = Grade.objects.get_or_create(
                student_id=student_id,
                subject_id=subject_id,
                sequence=sequence,
                academic_year_id=academic_year_id,
                evaluation_type=evaluation_type,
                defaults={
                    'teacher_id': g_data.get('teacher'),
                    'value': new_value
                }
            )
            
            if not is_new:
                try:
                    new_val_decimal = float(new_value)
                    old_val_decimal = float(grade_obj.value)
                except:
                    continue
                    
                if abs(new_val_decimal - old_val_decimal) > 0.01:
                    if not reason:
                        errors.append(f"Motif requis pour la modification de la note de {student.last_name}.")
                        continue
                        
                    GradeHistory.objects.create(
                        grade=grade_obj,
                        old_value=grade_obj.value,
                        new_value=new_value,
                        reason=reason,
                        changed_by=self.request.user
                    )
                    grade_obj.value = new_value
                    if g_data.get('teacher'):
                        grade_obj.teacher_id = g_data.get('teacher')
                    grade_obj.save()
                    updated += 1
            else:
                created += 1
                
        return Response({
            'message': 'Opération terminée.', 
            'created': created, 
            'updated': updated,
            'errors': errors
        })

    @action(detail=False, methods=['post'])
    def import_grades(self, request):
        if 'file' not in request.FILES:
            return Response({'error': 'Aucun fichier fourni.'}, status=400)
            
        school_class_id = request.data.get('school_class')
        subject_id = request.data.get('subject')
        sequence = request.data.get('sequence')
        academic_year_id = request.data.get('academic_year')
        evaluation_type = request.data.get('evaluation_type')
        teacher_id = request.data.get('teacher')
        
        if not all([school_class_id, subject_id, sequence, academic_year_id, evaluation_type]):
            return Response({'error': 'Veuillez fournir la classe, la matière, la séquence, l\'année académique et le type d\'évaluation.'}, status=400)
            
        file = request.FILES['file']
        
        created = 0
        updated = 0
        errors = []
        parsed_data = [] # list of dicts: {'identifier': str, 'note': float}
        
        try:
            if file.name.endswith('.xlsx'):
                import openpyxl
                wb = openpyxl.load_workbook(file, data_only=True)
                ws = wb.active
                for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                    if not row or not row[0]: continue
                    # Assuming format: Matricule/Nom | ... | Note
                    # We will try to find a float in the row, preferably at the end, and string at the beginning
                    identifier = str(row[0]).strip()
                    note = None
                    for cell in reversed(row):
                        if cell is not None:
                            try:
                                note = float(cell)
                                break
                            except:
                                pass
                    if note is not None:
                        parsed_data.append({'identifier': identifier, 'note': note, 'row': row_idx})
                    else:
                        errors.append(f"Ligne {row_idx} : Note invalide ou introuvable.")
            elif file.name.endswith('.docx'):
                from docx import Document
                doc = Document(file)
                if doc.tables:
                    table = doc.tables[0]
                    for row_idx, row in enumerate(table.rows):
                        if row_idx == 0: continue # header
                        if len(row.cells) >= 2:
                            identifier = row.cells[0].text.strip()
                            note_text = row.cells[-1].text.strip().replace(',', '.')
                            if not identifier: continue
                            try:
                                note = float(note_text)
                                parsed_data.append({'identifier': identifier, 'note': note, 'row': row_idx + 1})
                            except:
                                errors.append(f"Ligne {row_idx + 1} (Word) : Note '{note_text}' invalide.")
                else:
                    return Response({'error': 'Le fichier Word ne contient aucun tableau.'}, status=400)
            else:
                return Response({'error': 'Format non supporté. Utilisez .xlsx ou .docx'}, status=400)
                
            from platform_services.education.students.models import Student
            
            # Fetch all students in that class for the academic year
            students = Student.objects.filter(enrollments__school_class_id=school_class_id, enrollments__academic_year_id=academic_year_id).distinct()
            
            for item in parsed_data:
                identifier = item['identifier'].lower()
                note = item['note']
                if note < 0 or note > 20:
                    errors.append(f"Ligne {item['row']} : Note {note} hors limites (0-20).")
                    continue
                    
                # Match student
                matched_student = None
                for s in students:
                    if s.matricule and s.matricule.lower() == identifier:
                        matched_student = s
                        break
                    if s.last_name.lower() in identifier or s.first_name.lower() in identifier or identifier in f"{s.first_name} {s.last_name}".lower() or identifier in f"{s.last_name} {s.first_name}".lower():
                        matched_student = s
                        break
                        
                if not matched_student:
                    errors.append(f"Ligne {item['row']} : Élève '{item['identifier']}' introuvable dans la classe.")
                    continue
                    
                if self.check_is_locked(matched_student, academic_year_id, sequence):
                    errors.append(f"Ligne {item['row']} : Saisie verrouillée pour {matched_student.last_name}.")
                    continue
                    
                grade_obj, is_new = Grade.objects.get_or_create(
                    student=matched_student,
                    subject_id=subject_id,
                    sequence=sequence,
                    academic_year_id=academic_year_id,
                    evaluation_type=evaluation_type,
                    defaults={
                        'teacher_id': teacher_id or getattr(self.request.user, 'teacher_id', None),
                        'value': note
                    }
                )
                
                if not is_new:
                    if float(grade_obj.value) != note:
                        GradeHistory.objects.create(
                            grade=grade_obj,
                            old_value=grade_obj.value,
                            new_value=note,
                            reason="Importation fichier (mise à jour)",
                            changed_by=self.request.user
                        )
                        grade_obj.value = note
                        if teacher_id: grade_obj.teacher_id = teacher_id
                        grade_obj.save()
                        updated += 1
                else:
                    created += 1
                    
            return Response({
                'message': f'Import terminé. {created} créées, {updated} mises à jour.',
                'created': created,
                'updated': updated,
                'errors': errors
            })
        except Exception as e:
            return Response({'error': str(e)}, status=500)


class GradeHistoryViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = GradeHistory.objects.select_related('grade', 'changed_by').all()
    serializer_class = GradeHistorySerializer

    def get_queryset(self):
        queryset = super().get_queryset()
        student_id = self.request.query_params.get('student_id')
        class_id = self.request.query_params.get('class_id')
        academic_year_id = self.request.query_params.get('academic_year_id')
        sequence = self.request.query_params.get('sequence')

        if student_id:
            queryset = queryset.filter(grade__student_id=student_id)
        if class_id:
            # Need to filter through enrollments where school_class matches
            # And ideally for the specific academic year
            if academic_year_id:
                queryset = queryset.filter(
                    grade__student__enrollments__school_class_id=class_id,
                    grade__student__enrollments__academic_year_id=academic_year_id
                ).distinct()
            else:
                queryset = queryset.filter(grade__student__enrollments__school_class_id=class_id).distinct()
        if academic_year_id:
            queryset = queryset.filter(grade__academic_year_id=academic_year_id)
        if sequence:
            queryset = queryset.filter(grade__sequence=sequence)
            
        return queryset


class SequenceValidationViewSet(TenantQuerySetMixin, viewsets.ModelViewSet):
    queryset = SequenceValidation.objects.all()
    serializer_class = SequenceValidationSerializer

    @action(detail=False, methods=['post'])
    def toggle_lock(self, request):
        school_class_id = request.data.get('school_class')
        academic_year_id = request.data.get('academic_year')
        sequence = request.data.get('sequence')
        is_locked = request.data.get('is_locked', True)
        
        if not all([school_class_id, academic_year_id, sequence]):
            return Response({'error': 'Paramètres manquants.'}, status=400)
            
        validation, created = SequenceValidation.objects.get_or_create(
            school_class_id=school_class_id,
            academic_year_id=academic_year_id,
            sequence=sequence,
            defaults={'is_locked': is_locked, 'locked_by': request.user, 'locked_at': timezone.now()}
        )
        
        if not created:
            validation.is_locked = is_locked
            if is_locked:
                validation.locked_by = request.user
                validation.locked_at = timezone.now()
            validation.save()
            
        status_text = "verrouillée" if validation.is_locked else "déverrouillée"
        return Response({'message': f'Séquence {status_text} avec succès.', 'is_locked': validation.is_locked})
